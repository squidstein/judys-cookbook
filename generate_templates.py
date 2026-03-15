#!/usr/bin/env python3
"""
generate_templates.py — Bulk recipe template converter
=======================================================
Reads all old .doc recipe files and generates pre-filled .docx files
using the new recipe template format.

Judy can then open each .docx in Google Docs, fill in:
  - Description (a short blurb about the dish)
  - Chef's Notes (storage, substitutions, etc.)
  - Cuisine (if the script guessed wrong or left it blank)

Then save back to the iCloud Cookbook folder.

Usage:
    python3 generate_templates.py

Output:
    recipes_new/<Category>/RecipeName.docx
    (Does NOT overwrite existing .docx files — those are already done)
"""

import os, re, subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ──────────────────────────────────────────────────────────────────
RECIPES_DIR = os.path.join(os.path.dirname(__file__), "recipes")
OUTPUT_DIR  = os.path.join(os.path.dirname(__file__), "recipes_new")

# ── Kashrut detection ──────────────────────────────────────────────────────
MEAT_KEYWORDS = [
    'chicken', 'beef', 'lamb', 'veal', 'turkey', 'brisket', 'rib',
    'steak', 'meatball', 'meat loaf', 'meat ', 'short rib', 'burger',
    'sausage', 'duck', 'goose', 'liver', 'flanken', 'schnitzel',
    'kebab', 'deli meat', 'ground beef', 'ground turkey', 'ground lamb',
    'pastrami', 'salami', 'corned beef', 'cholent', 'pulpeta', 'carne',
]
DAIRY_KEYWORDS = [
    'milk', ' cream', 'cheese', 'butter', 'yogurt', 'sour cream',
    'ricotta', 'mozzarella', 'parmesan', 'cream cheese', 'half and half',
    'heavy cream', 'whipping cream', 'brie', 'cheddar', 'feta',
    'cottage cheese', 'buttermilk', 'whipped cream', 'ghee', 'kefir',
    'mascarpone', 'gruyere', 'swiss cheese', 'monterey',
]

def detect_kashrut(ingredients, steps):
    text = ' '.join(ingredients + steps).lower()
    has_meat  = any(kw in text for kw in MEAT_KEYWORDS)
    has_dairy = any(kw in text for kw in DAIRY_KEYWORDS)
    if has_meat and has_dairy:
        return 'Meat'   # conflict — flag as Meat, Judy can correct
    if has_meat:
        return 'Meat'
    if has_dairy:
        return 'Dairy'
    return 'Pareve'

# ── Cuisine guessing ───────────────────────────────────────────────────────
SEPHARDI_HINTS = [
    'adafina', 'boregas', 'borekas', 'moroccan', 'turkish', 'lentejas',
    'ropa vieja', 'carne guizado', 'potaje', 'oriza', 'pulpeta',
    'pastel de patata', 'ensalada de patata', 'ensalada cocha',
    'pescado', 'fijuela', 'biscocho', 'roskita', 'galleta de coco',
    'almond cigar', 'meat cigar', 'preserved lemon', 'somosa',
    'tagine', 'tangine', 'mango salsa', 'olive paste', 'eggplant',
    'carrot salad', 'beets', 'peppers', 'hummus', 'sababa',
    'spring vegetable', 'turkish salad', 'armenian lentil',
    'gypsy soup', 'tita sarah', 'quaziada', 'molde americana',
]
ASHKENAZI_HINTS = [
    'kugel', 'borscht', 'challah', 'latke', 'tzimmes', 'tsimmis',
    'cholent', 'rugelach', 'rugalach', 'matzo', 'gefilte', 'blintz',
    'kreplach', 'pecan keffl', 'forgotten cookie', 'gingerbread',
    'apple crisp', 'apple pie', 'apple cake', 'sour cream coffee',
    'lemon bundt', 'lemon poppy', 'lemon meringue', 'banana bread',
    'crescent roll', 'french bread', 'wheatgerm', 'cornmeal muffin',
    'beet borscht', 'pea soup', 'brisket', 'short ribs', 'meatball',
    'sweet meatball', 'meat loaf', 'fancy brisket', 'roasted brisket',
    'tiramisu', 'chocolate chip', 'oatmeal chip', 'toll house',
    'sour cream cookie', 'passover macaroon', 'round raisin',
]

def guess_cuisine(name, category):
    nl = name.lower()
    if any(h in nl for h in SEPHARDI_HINTS):
        return 'Sephardi'
    if any(h in nl for h in ASHKENAZI_HINTS):
        return 'Ashkenazi'
    # Category-level fallbacks
    if category == 'Passover':
        return 'Ashkenazi'
    return ''   # blank — Judy fills in

# ── Shared meta-line parser ────────────────────────────────────────────────
def extract_meta_fields(line, r):
    """
    Parse a single metadata line, handling old docs that cram multiple
    fields onto one line with tab separation, e.g.:
      "From: Grandma's Kitchen          Number of servings: 6-8"
    """
    chunks = [c.strip() for c in re.split(r' {2,}|\t', line) if c.strip()]
    for chunk in chunks:
        cl = chunk.lower()
        if cl.startswith("from"):
            val = re.split(r"from[:\s]+", chunk, flags=re.I, maxsplit=1)[-1].strip()
            if val and not r["from"]:
                r["from"] = val
        elif re.search(r"prep(?:aration)?\.?\s*time", cl):
            val = re.split(r"prep(?:aration)?\.?\s*time[:\s]+", chunk, flags=re.I, maxsplit=1)[-1].strip()
            if val and not r["prep_time"]:
                r["prep_time"] = val
        elif re.search(r"cook\.?\s*time", cl):
            val = re.split(r"cook\.?\s*time[:\s]+", chunk, flags=re.I, maxsplit=1)[-1].strip()
            if val and not r["cook_time"]:
                r["cook_time"] = val
        elif re.search(r"yield|serves|servings|number of", cl):
            val = re.split(r"(yields?|serves?|servings?|number of\s+servings?)[:\s]*", chunk, flags=re.I, maxsplit=1)[-1].strip()
            if val and not r["yields"]:
                r["yields"] = val
        else:
            r["meta"].append(chunk)


# ── Legacy .doc parser ─────────────────────────────────────────────────────
def read_doc(path):
    result = subprocess.run(
        ['textutil', '-convert', 'txt', '-stdout', path],
        capture_output=True, text=True
    )
    return result.stdout.strip()

def parse_legacy(name, text):
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    r = {
        'from': None, 'prep_time': None, 'cook_time': None,
        'yields': None, 'meta': [], 'ingredients': [], 'steps': [],
    }
    ing_idx = inst_idx = None
    for i, l in enumerate(lines):
        if re.match(r'^INGREDIENTS?\s*$', l, re.I): ing_idx = i
        if re.match(r'^INSTRUCTIONS?\s*$|^DIRECTIONS?\s*$|^METHOD\s*$|^PREPARATION\s*$', l, re.I): inst_idx = i

    meta_end = ing_idx if ing_idx is not None else (inst_idx if inst_idx is not None else len(lines))
    for l in lines[1:meta_end]:
        if l.lower() == name.lower() or len(l) > 150: continue
        extract_meta_fields(l, r)

    if ing_idx is not None:
        end = inst_idx if inst_idx is not None else len(lines)
        for l in lines[ing_idx+1:end]:
            if l and len(l) < 200:
                r['ingredients'].append(l)

    if inst_idx is not None:
        current = []
        for l in lines[inst_idx+1:]:
            if l:
                current.append(l)
            else:
                if current:
                    r['steps'].append(' '.join(current))
                    current = []
        if current:
            r['steps'].append(' '.join(current))

    return r

# ── .docx writer ───────────────────────────────────────────────────────────
def write_template(out_path, name, cuisine, from_, prep_time, cook_time,
                   yields, kashrut, ingredients, steps):
    doc = Document()

    # Title
    title = doc.add_paragraph(name)
    title.style = doc.styles['Title']

    # Metadata
    doc.add_paragraph(f'Cuisine: {cuisine}')
    doc.add_paragraph(f'From: {from_ or ""}')
    meta_block = (
        f'Prep Time: {prep_time or "-"}\n'
        f'Cook Time: {cook_time or "-"}\n'
        f'Yields: {yields or ""}\n'
        f'Kashrut: {kashrut}'
    )
    doc.add_paragraph(meta_block)

    # Description
    h = doc.add_paragraph('Description')
    h.style = doc.styles['Heading 1']
    doc.add_paragraph('')   # Judy fills this in

    # Ingredients
    h = doc.add_paragraph('Ingredients')
    h.style = doc.styles['Heading 1']
    for ing in ingredients:
        doc.add_paragraph(ing)

    # Preparation Steps
    h = doc.add_paragraph('Preparation Steps')
    h.style = doc.styles['Heading 1']
    for step in steps:
        doc.add_paragraph(step)

    # Chef's Notes
    h = doc.add_paragraph("Chef's Notes")
    h.style = doc.styles['Heading 1']
    doc.add_paragraph('Storage: ')
    doc.add_paragraph('Substitutions: ')
    doc.add_paragraph('Can you Freeze it: ')

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)

# ── Main ───────────────────────────────────────────────────────────────────
def main():
    SKIP = {
        'Cover Sheet', 'Doc1', 'Labels - Cookbook', 'Label for Spices',
        'Label for Spices 2', 'Side tabs', 'Spices - Aaron', 'Spices - Another',
        'Spices - Camila', 'Spices - Camila 1', 'Spices - Camila 2', 'Buttercup Icing',
    }
    DOC_EXTS = {'.doc', '.odt', '.rtf'}

    converted, skipped_existing, skipped_empty = 0, 0, 0
    cuisine_guesses = {'Sephardi': 0, 'Ashkenazi': 0, 'blank': 0}

    for cat in sorted(os.listdir(RECIPES_DIR)):
        cat_path = os.path.join(RECIPES_DIR, cat)
        if not os.path.isdir(cat_path) or cat.startswith('.'): continue

        # Track which recipes already have a .docx
        existing_docx = {
            os.path.splitext(f)[0]
            for f in os.listdir(cat_path) if f.endswith('.docx')
        }

        for fname in sorted(os.listdir(cat_path)):
            root, ext = os.path.splitext(fname)
            if ext.lower() not in DOC_EXTS: continue
            if root in SKIP: continue
            if root in existing_docx:
                skipped_existing += 1
                continue

            text = read_doc(os.path.join(cat_path, fname))
            if not text:
                skipped_empty += 1
                continue

            r = parse_legacy(root, text)
            cuisine  = guess_cuisine(root, cat)
            kashrut  = detect_kashrut(r['ingredients'], r['steps'])

            out_path = os.path.join(OUTPUT_DIR, cat, root + '.docx')
            write_template(
                out_path  = out_path,
                name      = root,
                cuisine   = cuisine,
                from_     = r['from'],
                prep_time = r['prep_time'],
                cook_time = r['cook_time'],
                yields    = r['yields'],
                kashrut   = kashrut,
                ingredients = r['ingredients'],
                steps     = r['steps'],
            )

            cuisine_guesses[cuisine if cuisine else 'blank'] += 1
            print(f'  ✓ {cat}/{root}  [{cuisine or "?cuisine"}, {kashrut}]')
            converted += 1

    print(f'\n── Summary ──────────────────────────────')
    print(f'  Converted:        {converted}')
    print(f'  Skipped (have .docx already): {skipped_existing}')
    print(f'  Skipped (empty):  {skipped_empty}')
    print(f'  Cuisine guesses:  Sephardi={cuisine_guesses["Sephardi"]}, '
          f'Ashkenazi={cuisine_guesses["Ashkenazi"]}, blank={cuisine_guesses["blank"]}')
    print(f'\nOutput folder: {OUTPUT_DIR}')
    print('Next step: review, then move .docx files into the recipes/ iCloud folder.')

if __name__ == '__main__':
    main()
